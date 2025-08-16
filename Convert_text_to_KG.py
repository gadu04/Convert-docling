from docx import Document
import re
import json
import uuid

# ===== CONFIG =====
DOCX_PATH = "Data/22.4.2024 UET Quy che chi tieu noi bo TV.docx"
OUTPUT_JSON = "output_schema.json"

# ===== HÀM TẠO ID =====
def gen_id(prefix):
    return f"{prefix}_{uuid.uuid4().hex[:8]}"

# ===== KHỞI TẠO CẤU TRÚC JSON =====
kg_data = {
    "Documents": [],
    "Chapters": [],
    "Sections": [],
    "Subsections": [],
    "Articles": [],
    "Procedures": [],
    "Tables": [],
    "Relationships": []
}

# ===== ĐỌC DOCX =====
doc = Document(DOCX_PATH)

# Tạo Document node
doc_id = gen_id("doc")
document_node = {
    "id": doc_id,
    "Name": DOCX_PATH.split("/")[-1].replace(".docx", ""),
    "Number_of_Pages": None,  # Không lấy được từ python-docx
    "Number_of_Chapter": 0
}
kg_data["Documents"].append(document_node)

# ===== BIẾN TẠM THEO DÕI =====
current_chapter = None
current_section = None
current_subsection = None
current_article = None
current_procedure = None
chapter_count = 0

# ===== DUYỆT PARAGRAPH =====
for para in doc.paragraphs:
    para_text = para.text.strip()
    if not para_text:
        continue

    # CHAPTER
    if re.match(r"^CHƯƠNG\s+\d+", para_text.upper()):
        chapter_count += 1
        chapter_id = gen_id("chapter")
        current_chapter = chapter_id
        kg_data["Chapters"].append({
            "id": chapter_id,
            "Name": para_text,
            "Number": chapter_count
        })
        kg_data["Relationships"].append({
            "from": doc_id, "to": chapter_id, "type": "HAS_CHAPTER"
        })
        document_node["Number_of_Chapter"] = chapter_count
        current_section = None
        current_subsection = None
        current_article = None
        current_procedure = None
        continue

    # SECTION
    if re.match(r"^Mục\s+\d+", para_text):
        section_id = gen_id("section")
        current_section = section_id
        kg_data["Sections"].append({
            "id": section_id,
            "Name": para_text,
            "Number": re.findall(r"\d+", para_text)[0]
        })
        if current_chapter:
            kg_data["Relationships"].append({
                "from": current_chapter, "to": section_id, "type": "HAS_SECTION"
            })
        current_subsection = None
        current_article = None
        current_procedure = None
        continue

    # SUBSECTION
    if re.match(r"^[A-Z]\.", para_text):
        sub_id = gen_id("subsection")
        current_subsection = sub_id
        kg_data["Subsections"].append({
            "id": sub_id,
            "Name": para_text,
            "Letter": para_text.split(".")[0]
        })
        if current_section:
            kg_data["Relationships"].append({
                "from": current_section, "to": sub_id, "type": "CONTAINS"
            })
        continue

    # ARTICLE
    if re.match(r"^Điều\s+\d+", para_text):
        art_id = gen_id("article")
        current_article = art_id
        kg_data["Articles"].append({
            "id": art_id,
            "Name": para_text,
            "Number": re.findall(r"\d+", para_text)[0]
        })
        if current_section:
            kg_data["Relationships"].append({
                "from": current_section, "to": art_id, "type": "HAS_ARTICLE"
            })
        continue

    # PROCEDURE
    if re.match(r"^Thủ tục\s+\d+", para_text):
        proc_id = gen_id("procedure")
        current_procedure = proc_id
        kg_data["Procedures"].append({
            "id": proc_id,
            "Name": para_text,
            "Number": re.findall(r"\d+", para_text)[0]
        })
        if current_section:
            kg_data["Relationships"].append({
                "from": current_section, "to": proc_id, "type": "HAS_PROCEDURE"
            })
        continue

# ===== DUYỆT BẢNG =====
for table in doc.tables:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    table_id = gen_id("table")
    kg_data["Tables"].append({
        "id": table_id,
        "Name": f"Table_{len(kg_data['Tables'])+1}",
        "Reference": rows
    })

    # Gắn bảng vào Article hoặc Procedure gần nhất
    if current_article:
        kg_data["Relationships"].append({
            "from": current_article, "to": table_id, "type": "HAS_TABLE"
        })
    elif current_procedure:
        kg_data["Relationships"].append({
            "from": current_procedure, "to": table_id, "type": "HAS_TABLE"
        })

# ===== LƯU RA JSON =====
with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
    json.dump(kg_data, f, ensure_ascii=False, indent=2)

print(f"✅ Đã xuất JSON theo schema vào {OUTPUT_JSON}")
